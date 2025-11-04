from pathlib import Path
import datetime
import pytest

from expedienteindex.exporters import export_docx, export_pdf

try:
    from docx import Document as _DocxReader
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False
    _DocxReader = None
    _WD_ALIGN = None

try:
    from pypdf import PdfReader as _PdfReader
    PYPDF_AVAILABLE = True
except Exception:
    PYPDF_AVAILABLE = False
    _PdfReader = None


def _read_docx_texts(doc_path: Path):
    """Returns a list with each paragraph text from the docx."""
    doc = _DocxReader(str(doc_path))
    return [p.text for p in doc.paragraphs]

@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not available")
def test_docx_defaults_show_title_and_date(tmp_path: Path):
    titles = ["A", "B"]
    out = tmp_path / "idx.docx"
    export_docx(titles, out)
    assert out.exists() and out.stat().st_size > 0

    texts = _read_docx_texts(out)
    today = datetime.date.today().strftime("%d/%m/%Y")

    assert any(t == "Índice de Documentos" for t in texts)
    assert any(today in t for t in texts)
    assert any("- A" in t for t in texts)
    assert any("- B" in t for t in texts)

@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not available")
def test_docx_custom_title_no_date_and_alignment_right(tmp_path: Path):
    titles = ["One", "Two"]
    out = tmp_path / "idx_custom.docx"
    custom_title = "Índice del Procedimiento 123/2025"

    export_docx(
        titles, out,
        title_text=custom_title,
        show_title=True,
        show_date=False,
        title_align="right"
    )
    assert out.exists() and out.stat().st_size > 0

    doc = _DocxReader(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert len(texts) > 0 and texts[0] == custom_title
    assert doc.paragraphs[0].alignment == _WD_ALIGN.RIGHT
    today = datetime.date.today().strftime("%d/%m/%Y")
    assert all(today not in t for t in texts)

@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not available")
def test_docx_hide_title_only_date_visible(tmp_path: Path):
    titles = ["DocX"]
    out = tmp_path / "idx_only_date.docx"

    export_docx(
        titles, out,
        show_title=False,
        show_date=True,
        title_align="center"
    )
    assert out.exists() and out.stat().st_size > 0

    doc = _DocxReader(str(out))
    texts = [p.text for p in doc.paragraphs]
    today = datetime.date.today().strftime("%d/%m/%Y")

    assert all("Índice de Documentos" not in t for t in texts)
    assert any(today in t for t in texts)
    assert any("- DocX" in t for t in texts)

@pytest.mark.skipif(not PYPDF_AVAILABLE, reason="pypdf not available")
def test_pdf_defaults_show_title_and_date(tmp_path: Path):
    titles = ["A", "B"]
    out = tmp_path / "idx.pdf"
    export_pdf(titles, out)
    assert out.exists() and out.stat().st_size > 0

    reader = _PdfReader(str(out))
    text = "".join((page.extract_text() or "") for page in reader.pages)
    assert "Índice de Documentos" in text
    assert datetime.date.today().strftime("%d/%m/%Y") in text
    assert "- A" in text or "A" in text # some extractors omit hyphen
    assert "- B" in text or "B" in text

@pytest.mark.skipif(not PYPDF_AVAILABLE, reason="pypdf not available")
def test_pdf_custom_title_no_date(tmp_path: Path):
    titles = ["One", "Two"]
    out = tmp_path / "idx_custom.pdf"
    custom_title = "Índice del Procedimiento 123/2025"

    export_pdf(
        titles, out,
        title_text=custom_title,
        show_title=True,
        show_date=False,
        title_align="left"
    )
    assert out.exists() and out.stat().st_size > 0

    reader = _PdfReader(str(out))
    text = "".join((page.extract_text() or "") for page in reader.pages)
    assert custom_title in text
    assert datetime.date.today().strftime("%d/%m/%Y") not in text