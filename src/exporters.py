import datetime
from pathlib import Path
from typing import List

# ---- DOCX ----
def export_docx(titles: List[str], out_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    h = doc.add_paragraph()
    run = h.add_run("Índice de Documentos")
    run.bold = True
    run.font.size = Pt(18)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_p = doc.add_paragraph()
    date_run = date_p.add_run(datetime.date.today().strftime("%d/%m/%Y"))
    date_run.italic = True
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    for t in titles:
        para = doc.add_paragraph(f"- {t}")
        para.paragraph_format.space_after = Pt(4)

    doc.save(out_path)

# ---- PDF ----
def export_pdf(titles: List[str], out_path: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import cm

    c = canvas.Canvas(str(out_path), pagesize=A4)
    width, height = A4

    top_margin = 2.5 * cm
    left_margin = 2.5 * cm
    y = height - top_margin

    c.setFont("Helvetica-Bold", 18)
    c.drawCentredString(width / 2, y, "Índice de Documentos")
    y -= 0.9 * cm

    c.setFont("Helvetica-Oblique", 10)
    c.drawCentredString(width / 2, y, datetime.date.today().strftime("%d/%m/%Y"))
    y -= 1.0 * cm

    c.setFont("Helvetica", 11)
    line_height = 0.6 * cm

    for t in titles:
        if y < 2.5 * cm:
            c.showPage()
            y = height - top_margin
            c.setFont("Helvetica", 11)
        c.drawString(left_margin, y, f"- {t}")
        y -= line_height

    c.showPage()
    c.save()