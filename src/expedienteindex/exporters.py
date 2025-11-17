import datetime
from pathlib import Path
from typing import List, Optional

# ---- Helpers PDF font registration ----
_BASE14 = {"Helvetica", "Times-Roman", "Courier"}

def _register_pdf_font_id_needed(font_name: str) -> str:
    """
    Try to register TTF for ReportLab based on matplotlib's font manager.
    Returns the usable font name dor canvas.setFont().
    Falls back to Helvetica if not found/registrable.
    """
    try:
        from matplotlib import font_manager
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        if font_name in _BASE14:
            return font_name
        
        # FInd a matching font file by family name
        matches = [f for f in font_manager.fontManager.ttflist
                   if f.name.lower() == font_name.lower()]
        if not matches:
            return "Helvetica"
        
        font_path = matches[0].fname
        # Register under the family name
        if font_name not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(font_name, font_path))
        return font_name
    except Exception:
        return "Helvetica"

# ---- DOCX ----
def export_docx(
    titles: List[str], 
    out_path: Path,
    *,
    title_text: str = "Índice de Documentos",
    show_title: bool = True,
    show_date: bool = True,
    title_align: str = "center", # left / center / right    
    font_name: str = "Calibri",
    title_font_size: int = 18,
    body_font_size: int = 11,
) -> None:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }

    if show_title:
        h = doc.add_paragraph()
        run = h.add_run(title_text)
        run.bold = True
        run.font.size = Pt(int(title_font_size))
        run.font.name = font_name
        h.alignment = align_map.get(title_align, WD_ALIGN_PARAGRAPH.CENTER)
        
    if show_date:
        date_p = doc.add_paragraph()
        date_run = date_p.add_run(datetime.date.today().strftime("%d/%m/%Y"))
        date_run.italic = True
        date_run.font.name = font_name
        date_run.font.size = Pt(max(9, int(body_font_size) - 1))
        date_p.alignment = align_map.get(title_align, WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    for t in titles:
        para = doc.add_paragraph(f"{t}")
        para.paragraph_format.space_after = Pt(4)
        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(int(body_font_size))

    doc.save(out_path)

# ---- PDF ----
def export_pdf(
    titles: List[str], 
    out_path: Path,
    *,
    title_text: str = "Índice de Documentos",
    show_title: bool = True,
    show_date: bool = True,
    title_align: str = "center",
    font_name: str = "Helvetica",
    title_font_size: int = 18,
    body_font_size: int = 11,
) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import cm

    c = canvas.Canvas(str(out_path), pagesize=A4)
    width, height = A4

    top_margin = 2.5 * cm
    left_margin = 2.5 * cm
    y = height - top_margin

    usable_font = _register_pdf_font_id_needed(font_name)

    def draw_title_line(text: str, size: int = 18):
        c.setFont(font_name, int(size))
        if title_align == "left":
            c.drawString(left_margin, y, text)
        elif title_align == "right":
            c.drawRightString(width - left_margin, y, text)
        else:
            c.drawCentredString(width / 2, y, text)

    def draw_date_line(text: str, size: int = 10):
        c.setFont(usable_font, int(size))
        if title_align == "left":
            c.drawString(left_margin, y, text)
        elif title_align == "right":
            c.drawRightString(width - left_margin, y, text)
        else:
            c.drawCentredString(width / 2, y, text)
    
    if show_title:
        draw_title_line(title_text, title_font_size)
        y -= 0.9 * cm

    if show_date:
        draw_date_line(datetime.date.today().strftime("%d/%m/%Y"), max(9, int(body_font_size) - 1))
        y -= 1.0 * cm

    c.setFont(font_name, body_font_size)
    line_height = 0.6 * cm
    for t in titles:
        if y < 2.5 * cm:
            c.showPage()
            y = height - top_margin
            c.setFont(font_name, body_font_size)
        c.drawString(left_margin, y, f"{t}")
        y -= line_height

    c.showPage()
    c.save()